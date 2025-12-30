# save as make_docx.py
from docx import Document
from docx.shared import Pt

doc = Document()
title = doc.add_heading("AutoCV - Project Overview", 0)
title.runs[0].font.size = Pt(24)

doc.add_heading("Purpose", level=1)
doc.add_paragraph(
    "Accept resumes (PDF/DOCX), parse, score against target roles, and return a report with prioritized feedback."
)

doc.add_heading("Tech Stack", level=1)
doc.add_paragraph(
    "Backend: Flask 3.x (Python), SQLAlchemy/SQLite (swap to Postgres-ready), Gunicorn for prod\n"
    "Parsing: PyMuPDF (fitz) for PDF, python-docx for DOCX\n"
    "NLP/Scoring: spaCy en_core_web_sm, sentence-transformers, custom scoring logic\n"
    "Frontend: Jinja templates with Tailwind-like utility classes\n"
    "Config: config.py + .env (SECRET_KEY), python-dotenv\n"
)

doc.add_heading("Flow", level=1)
flow = doc.add_paragraph()
steps = [
    "Upload: Flask receives resume at /upload, saves to uploads/.",
    "Parse: core/parser.py uses fitz for PDF or python-docx for DOCX; cleans text/sections.",
    "Skills/Sections: core/skills.py loads data/skills_taxonomy.json; core/sections.py maps sections.",
    "Score: core/scorer.py & core/matcher.py compute sub-scores and overall score.",
    "Feedback: core/feedback.py & core/ats.py produce High/Medium/Low priorities and bullet rewrites.",
    "Persist: SQLAlchemy saves report/feedback; SQLite by default.",
    "Render: /report/<id> uses templates/report.html to display scores and feedback.",
]
for s in steps:
    flow.add_run(f"• {s}\n")

doc.add_heading("Run Locally", level=1)
doc.add_paragraph(
    "python -m venv .venv\n"
    "source .venv/bin/activate\n"
    "pip install --upgrade pip setuptools wheel\n"
    "pip install -r requirements.txt\n"
    "echo \"SECRET_KEY=replace-me\" > .env\n"
    "python app.py  # visits http://127.0.0.1:5000\n"
)

doc.add_heading("Deployment Notes", level=1)
doc.add_paragraph(
    "Use Python 3.11.x (PyMuPDF wheels). Start cmd: gunicorn app:app --workers 2 --bind 0.0.0.0:$PORT.\n"
    "Set SECRET_KEY and DATABASE_URL in env; default is SQLite."
)

doc.save("project_overview.docx")
print("Wrote project_overview.docx")